"""
Unit tests for Document Extraction.
"""
import os
import tempfile
import docx
from app.services.document_extractor import extract_text_from_file, extract_text_from_txt, extract_text_from_docx

def test_extract_txt():
    sample_text = "This is a sample academic paper on distributed computing and cloud security."
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8") as f:
        f.write(sample_text)
        temp_path = f.name

    try:
        extracted, w_count, c_count = extract_text_from_file(temp_path, "sample.txt")
        assert sample_text in extracted
        assert w_count == 12
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def test_extract_docx():
    doc = docx.Document()
    doc.add_paragraph("Machine learning algorithms evaluate security telemetry.")
    doc.add_paragraph("Second paragraph regarding feature selection and entropy.")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name
        doc.save(temp_path)

    try:
        extracted, w_count, c_count = extract_text_from_file(temp_path, "sample.docx")
        assert "Machine learning algorithms" in extracted
        assert "entropy" in extracted
        assert w_count > 10
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
