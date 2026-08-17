"""
Document Text Extraction Service for PDF, DOCX, and TXT files.
"""
import os
import re
from typing import Tuple

try:
    import pymupdf
    HAVE_PYMUPDF = True
except ImportError:
    HAVE_PYMUPDF = False

try:
    import pypdf
    HAVE_PYPDF = True
except ImportError:
    HAVE_PYPDF = False

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


class DocumentExtractionError(Exception):
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz) or pypdf."""
    text_content = []

    if HAVE_PYMUPDF:
        try:
            doc = pymupdf.open(file_path)
            for page in doc:
                text_content.append(page.get_text())
            doc.close()
            extracted = "\n".join(text_content).strip()
            if extracted:
                return extracted
        except Exception:
            pass # Try fallback if PyMuPDF fails

    if HAVE_PYPDF:
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                extracted_page = page.extract_text()
                if extracted_page:
                    text_content.append(extracted_page)
            return "\n".join(text_content).strip()
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract PDF using pypdf: {str(e)}")

    if not text_content:
        raise DocumentExtractionError("No valid PDF extraction library available or document is empty/scanned.")
    return "\n".join(text_content).strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    if not HAVE_DOCX:
        raise DocumentExtractionError("python-docx library is not installed.")
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as e:
        raise DocumentExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file using UTF-8 with fallback encodings."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise DocumentExtractionError(f"Error reading TXT file: {str(e)}")
    raise DocumentExtractionError("Unable to decode text file with standard encodings.")


def extract_text_from_file(file_path: str, filename: str) -> Tuple[str, int, int]:
    """
    Extracts text from a document based on its extension.
    Returns: (extracted_text, word_count, character_count)
    """
    if not os.path.exists(file_path):
        raise DocumentExtractionError(f"File not found at path: {file_path}")

    _, ext = os.path.splitext(filename.lower())

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in [".docx"]:
        text = extract_text_from_docx(file_path)
    elif ext in [".txt", ".text", ".md"]:
        text = extract_text_from_txt(file_path)
    elif ext == ".doc":
        # Legacy binary .doc format fallback
        try:
            text = extract_text_from_txt(file_path)
        except Exception:
            raise DocumentExtractionError(
                "Legacy .doc format detected. Please save your file as modern .docx or .pdf for accurate text analysis."
            )
    else:
        raise DocumentExtractionError(f"Unsupported file format: {ext}. Allowed: PDF, DOCX, TXT.")

    # Clean non-printable control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    
    words = [w for w in text.split() if w]
    word_count = len(words)
    char_count = len(text)

    return text, word_count, char_count
